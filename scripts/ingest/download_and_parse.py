import csv
import json
import re
import sys
import time
import urllib.request
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent.parent
DATA = ROOT / "data" / "statystyka-policja"
MANIFEST = json.loads((ROOT / "scripts" / "manifest.json").read_text(encoding="utf-8"))
PROGRESS_FILE = DATA / "progress.json"

HEADERS = {"User-Agent": "Mozilla/5.0 (research data collection; polistats project)"}


def load_progress() -> dict:
    if PROGRESS_FILE.exists():
        return json.loads(PROGRESS_FILE.read_text(encoding="utf-8"))
    return {}


def save_progress(progress: dict):
    DATA.mkdir(parents=True, exist_ok=True)
    PROGRESS_FILE.write_text(json.dumps(progress, ensure_ascii=False, indent=2), encoding="utf-8")


def folder_name(filename: str) -> str:
    if "." in filename:
        name, ext = filename.rsplit(".", 1)
        return f"{name}_{ext}"
    return filename


def download(url: str, dest: Path) -> bool:
    if dest.exists() and dest.stat().st_size > 0:
        return True
    req = urllib.request.Request(url, headers=HEADERS)
    last_err = None
    for attempt in range(3):
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                data = resp.read()
            dest.write_bytes(data)
            return True
        except Exception as e:
            last_err = e
            wait = min(15, 3 * (2 ** attempt))
            print(f"  ! attempt {attempt+1} failed ({e}), retrying in {wait}s", file=sys.stderr, flush=True)
            time.sleep(wait)
    print(f"  ! download failed after retries: {last_err}", file=sys.stderr, flush=True)
    return False


def write_metadata(folder: Path, entry: dict, size_bytes: int):
    md = f"""# {entry['filename']}

- **Label / description:** {entry['label']}
- **Download link:** {entry['url']}
- **Source page (subsite):** {entry['page_url'] or '(see site-map)'}
- **Site-map reference:** site-map/{entry['source_md']}
- **Category:** {entry['category']}
- **File size:** {size_bytes:,} bytes
"""
    (folder / "metadata.md").write_text(md, encoding="utf-8")


def _load_workbook_sanitized(path: Path):
    """Some workbooks (often re-saved by WPS/Kingsoft Office) embed vendor
    <extLst> extensions inside style elements that openpyxl's strict
    PatternFill parser rejects. Strip them from styles.xml and retry --
    they're cosmetic font/fill metadata, not data.
    """
    import io
    import re
    import zipfile

    from openpyxl import load_workbook

    try:
        return load_workbook(path, data_only=True, read_only=True)
    except TypeError:
        pass

    buf = io.BytesIO()
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "xl/styles.xml":
                text = re.sub(rb"<extLst>.*?</extLst>", b"", data, flags=re.DOTALL)
                data = text
            zout.writestr(item, data)
    buf.seek(0)
    return load_workbook(buf, data_only=True, read_only=True)


def parse_xlsx(path: Path, tables_dir: Path) -> str:
    wb = _load_workbook_sanitized(path)
    lines = [f"Workbook with {len(wb.sheetnames)} sheet(s): {', '.join(wb.sheetnames)}\n"]
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        lines.append(f"\n## Sheet: {sheet_name} ({len(rows)} rows)\n")
        if not rows:
            continue
        tables_dir.mkdir(exist_ok=True)
        csv_path = tables_dir / f"{slug(sheet_name)}.csv"
        with csv_path.open("w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            for row in rows:
                w.writerow(["" if v is None else v for v in row])
        preview = rows[:15]
        for row in preview:
            lines.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
        if len(rows) > 15:
            lines.append(f"... ({len(rows) - 15} more rows, see tables/{csv_path.name})")
    return "\n".join(lines)


def parse_xls(path: Path, tables_dir: Path) -> str:
    import xlrd

    wb = xlrd.open_workbook(str(path))
    lines = [f"Workbook with {wb.nsheets} sheet(s): {', '.join(wb.sheet_names())}\n"]
    for sheet_name in wb.sheet_names():
        sh = wb.sheet_by_name(sheet_name)
        lines.append(f"\n## Sheet: {sheet_name} ({sh.nrows} rows)\n")
        if sh.nrows == 0:
            continue
        tables_dir.mkdir(exist_ok=True)
        csv_path = tables_dir / f"{slug(sheet_name)}.csv"
        with csv_path.open("w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            for r in range(sh.nrows):
                w.writerow(sh.row_values(r))
        for r in range(min(sh.nrows, 15)):
            lines.append("| " + " | ".join(str(v) for v in sh.row_values(r)) + " |")
        if sh.nrows > 15:
            lines.append(f"... ({sh.nrows - 15} more rows, see tables/{csv_path.name})")
    return "\n".join(lines)


def parse_pdf(path: Path, tables_dir: Path) -> str:
    import pdfplumber
    from pdf_tables import extract_tables_from_page, find_label_value_lists

    lines = []
    table_count = 0
    with pdfplumber.open(str(path)) as pdf:
        lines.append(f"PDF with {len(pdf.pages)} page(s).\n")
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                lines.append(f"\n## Page {i} text\n\n{text.strip()}\n")
            try:
                tables = extract_tables_from_page(page)
            except Exception as e:
                tables = []
                lines.append(f"\n(table extraction failed on page {i}: {e})\n")
            for t in tables:
                if not t or len(t) < 2:
                    continue
                table_count += 1
                tables_dir.mkdir(exist_ok=True)
                csv_path = tables_dir / f"page{i}-table{table_count}.csv"
                with csv_path.open("w", newline="", encoding="utf-8") as f:
                    w = csv.writer(f)
                    for row in t:
                        w.writerow(["" if v is None else v for v in row])
                lines.append(f"\n(Table extracted from page {i} -> tables/{csv_path.name}, {len(t)} rows)\n")

            # Only look for narrative bullet lists where no ruled table was
            # found -- a page with a real grid table sometimes linearizes
            # into text that coincidentally also matches "label - number".
            if tables:
                continue
            for heading, items in find_label_value_lists(text):
                table_count += 1
                tables_dir.mkdir(exist_ok=True)
                csv_path = tables_dir / f"page{i}-table{table_count}.csv"
                with csv_path.open("w", newline="", encoding="utf-8") as f:
                    w = csv.writer(f)
                    w.writerow([heading or "label", "value"])
                    w.writerows(items)
                label = f' "{heading}"' if heading else ""
                lines.append(f"\n(List{label} extracted from page {i} -> tables/{csv_path.name}, {len(items)} items)\n")
    return "\n".join(lines)


def parse_docx(path: Path, tables_dir: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for ti, table in enumerate(doc.tables, start=1):
        tables_dir.mkdir(exist_ok=True)
        csv_path = tables_dir / f"table{ti}.csv"
        with csv_path.open("w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            for row in table.rows:
                w.writerow([c.text for c in row.cells])
        lines.append(f"\n(Table {ti} extracted -> tables/{csv_path.name}, {len(table.rows)} rows)\n")
    return "\n".join(lines) if lines else "(no extractable text/tables — may be a legacy .doc binary format)"


def slug(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_\-]+", "_", s.strip())
    return s or "sheet"


def main():
    only_index = None
    if len(sys.argv) > 2 and sys.argv[1] == "--range":
        a, b = sys.argv[2].split(":")
        only_index = (int(a), int(b))

    entries = MANIFEST
    if only_index:
        entries = entries[only_index[0]:only_index[1]]

    progress = load_progress()

    for n, entry in enumerate(entries):
        key = f"{entry['category']}/{entry['filename']}"
        cat_dir = DATA / entry["category"]
        fdir = cat_dir / folder_name(entry["filename"])
        fdir.mkdir(parents=True, exist_ok=True)
        dest = fdir / entry["filename"]
        print(f"[{n+1}/{len(entries)}] {key}", flush=True)

        if progress.get(key, {}).get("status") == "done":
            continue

        err_file = fdir / "ERROR.txt"
        ok = download(entry["url"], dest)
        if not ok:
            err_file.write_text("download failed\n", encoding="utf-8")
            progress[key] = {"status": "error", "stage": "download"}
            save_progress(progress)
            continue
        if err_file.exists():
            err_file.unlink()

        write_metadata(fdir, entry, dest.stat().st_size)

        ext = entry["filename"].rsplit(".", 1)[-1].lower()
        tables_dir = fdir / "tables"
        content = None
        try:
            if ext == "xlsx":
                content = parse_xlsx(dest, tables_dir)
            elif ext == "xls":
                content = parse_xls(dest, tables_dir)
            elif ext == "pdf":
                content = parse_pdf(dest, tables_dir)
            elif ext == "docx":
                content = parse_docx(dest, tables_dir)
            elif ext == "doc":
                content = "(legacy .doc binary format — not parsed; open manually or convert with libreoffice)"
            else:
                content = f"(unhandled extension: {ext})"
        except Exception as e:
            content = f"(parsing failed: {e})"

        (fdir / "content.md").write_text(
            f"# Parsed content: {entry['filename']}\n\n{content}\n", encoding="utf-8"
        )
        progress[key] = {"status": "done"}
        save_progress(progress)
        time.sleep(2.0)


if __name__ == "__main__":
    main()
